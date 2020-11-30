from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from datetime import datetime, date

import pandas as pd
from copy import deepcopy

###############################################################################
import geopandas
import contextily as ctx
from shapely import wkt
###############################################################################

import matplotlib.pyplot as plt
import matplotlib.ticker as ticker


from ST_AM_Contingencies_Ploty import Radar_Plot_by_Asset, plot_historic_condition_by_asset, f_risk_aversion, f_Cr_Hist, f_Risk_Matrix, f_Pareto,f_Pareto_Year
#from Processing_tools import Report_df
from Processing_tools import Report_APM_df as Report_df 


file_name = 'RESULTS/X1_REPORTS/X1_Images/fig.png'

def Test_Report_AP(data,assets,plan_horizonts=None):
    document = Document('STATIC/APM_General_Report.docx')
    
    document.paragraphs[0].text = data['Name']
    document.paragraphs[1].text = data['Sub_title']
    report_date                 = datetime.date(datetime.now())
    
    if plan_horizonts==None:
        plan_horizonts = [report_date]
    else:
        plan_horizonts = [date(x, 12, 31) for x in plan_horizonts]
        print(plan_horizonts)

    document.paragraphs[2].text = str(report_date)

    df = Report_df(assets,report_date)

    #print(df)
    # # # # # # # # # # # # # # #
    document.tables[0]._cells[0].paragraphs[0].text = str(df['Name'].count())
    document.tables[0]._cells[1].paragraphs[0].text = str(round(datetime.now().year-df['Year'].mean(),1))
    document.tables[0]._cells[2].paragraphs[0].text = str(round(df['HI'].mean(),2))
    document.tables[0]._cells[3].paragraphs[0].text = str(round(df['RE'].mean(),2))
    # # # # # # # # # # # # # # #
    
    def l_save_figure(table,cells,p_type='Plot'):
        if p_type=='Plot':
            plt.ylabel('Número de activos',fontsize= 14)
            plt.xticks(fontsize=12)
            plt.yticks(fontsize=12)
            ax.yaxis.set_major_formatter(ticker.FormatStrFormatter("%d"))
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            #ax.grid(color='grey', linestyle='-', linewidth=0.25, alpha=0.5)

        plt.savefig(file_name, bbox_inches='tight')
        plt.close()
        paragraph = document.tables[table]._cells[cells].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(file_name, height=Inches(2.4))
    
    # Hi plot
    ax = df['HI'].plot.hist(alpha=1,color='#e81123')
    l_save_figure(1,2)
    # EL plot 
    ax =df['EL'].plot.hist(alpha=1,color='#ff8c00')
    l_save_figure(2,2)
    # Operative life plot
    ax =df['Year'].plot.hist(alpha=1,color='#009e49')
    ax.xaxis.set_major_formatter(ticker.FormatStrFormatter("%d"))
    l_save_figure(2,3)
    # Asset location plot
    ax = df.plot(figsize=(10, 8), alpha=0.75, edgecolor='k',color='#00188f')
    ctx.add_basemap(ax, zoom=10)
    ax.set_axis_off()
    l_save_figure(1,3,p_type='map')

    
    # Each asset report
    n_table = 3
    for asset_id in  list(assets.Asset_Portfolio_List.index):
        # Add title 
        asset     = assets.Asset_Portfolio[asset_id]
        paragraph = document.add_paragraph(asset.name)
        paragraph.style = document.styles['Title']
        # Table with indexes 
        template = document.tables[0]
        tbl      = template._tbl
        new_tbl = deepcopy(tbl)
        paragraph._p.addnext(new_tbl)

        document.tables[n_table]._cells[0].paragraphs[0].text = asset.type
        document.tables[n_table]._cells[0].paragraphs[1].text = 'Tipo de activo'

        document.tables[n_table]._cells[1].paragraphs[0].text = asset.data['Location']
        document.tables[n_table]._cells[1].paragraphs[1].text = 'Subestación'

        document.tables[n_table]._cells[2].paragraphs[0].text = str(asset.data['Serial'])
        document.tables[n_table]._cells[2].paragraphs[1].text = 'Serial'

        document.tables[n_table]._cells[3].paragraphs[0].text = asset.data['UC']
        document.tables[n_table]._cells[3].paragraphs[1].text = 'UC'

        n_table +=1 
# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #  
        paragraph = document.add_paragraph()  

        #file_name = 'RESULTS/X1_REPORTS/X1_Images/fig.png'
        plot_historic_condition_by_asset(asset,file_name)
        run = paragraph.add_run()
        run.add_picture(file_name, height=Inches(3))

# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #
        new_tbl = deepcopy(tbl)
        paragraph._p.addnext(new_tbl)
        document.tables[n_table]._cells[3].paragraphs[0].text = str(round(asset.mttr,2))
        document.tables[n_table]._cells[3].paragraphs[1].text = 'MTTR - [h]'

        document.tables[n_table]._cells[1].paragraphs[0].text = str(round(asset.elap_life,2))
        document.tables[n_table]._cells[1].paragraphs[1].text = 'EL'

        if asset.hi == None:
            document.tables[n_table]._cells[0].paragraphs[0].text = 'None'
        else:
            document.tables[n_table]._cells[0].paragraphs[0].text = str(round(asset.hi,2))
        document.tables[n_table]._cells[0].paragraphs[1].text = 'HI '

        if asset.re==None:
            document.tables[n_table]._cells[2].paragraphs[0].text = 'None'
        else:
            document.tables[n_table]._cells[2].paragraphs[0].text = str(round(asset.re,2))
        document.tables[n_table]._cells[2].paragraphs[1].text = 'RE'
        n_table +=1 
# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #  
        paragraph = document.add_paragraph()  
        #file_name = 'RESULTS/X1_REPORTS/X1_Images/fig.png'
        #Radar_Plot_by_Asset(asset,[report_date],file_name)
        Radar_Plot_by_Asset(asset,plan_horizonts,file_name)
        run = paragraph.add_run()
        run.add_picture(file_name, height=Inches(3.5))

# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #
        new_tbl = deepcopy(tbl)
        paragraph._p.addnext(new_tbl)

        document.tables[n_table]._cells[0].paragraphs[0].text = asset.data['ID']
        document.tables[n_table]._cells[0].paragraphs[1].text = 'ID'

        document.tables[n_table]._cells[1].paragraphs[0].text = str(asset.data['Opt_Year'].year)
        document.tables[n_table]._cells[1].paragraphs[1].text = 'Año de ent. operación'

        document.tables[n_table]._cells[2].paragraphs[0].text = asset.data['Brand']
        document.tables[n_table]._cells[2].paragraphs[1].text = 'Marca'

        document.tables[n_table]._cells[3].paragraphs[0].text = asset.data['Model']
        document.tables[n_table]._cells[3].paragraphs[1].text = 'Modelo'
        n_table +=1 

        #document.add_page_break()
        
    #Save document 
    document.save('RESULTS/X1_REPORTS/CONDITION.docx')

# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #
#                                                           #
#                                                           #     
#                                                           #    
#                                                           #   
# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #

def Test_Report_AC(data,df_mont,df,Years,N):
    document = Document('STATIC/ACM_General_Report.docx')
    
    document.paragraphs[0].text = data['Name']
    document.paragraphs[1].text = data['Sub_title']
    report_date                 = datetime.date(datetime.now())
    document.paragraphs[2].text = str(report_date)

    # # # # # # # # # # # # # # #
    document.tables[0]._cells[0].paragraphs[0].text = str(round(df[df['Year']==Years[0]]['A'].mean(),3))
    document.tables[0]._cells[1].paragraphs[0].text = str(round(df[df['Year']==Years[0]]['ENS'].sum(),2))
    document.tables[0]._cells[2].paragraphs[0].text = str(round(df[df['Year']==Years[0]]['SAIDI'].sum(),3))
    document.tables[0]._cells[3].paragraphs[0].text = str(round(df[df['Year']==Years[0]]['RI'].sum(),2))
    # # # # # # # # # # # # # # #

    from Processing_tools import CR_Time_df

    df_cr_time = CR_Time_df(df_mont,Years)

    # Plot aversion risk
    f_risk_aversion(df_cr_time,Years,file_name)
    paragraph   = document.add_paragraph() 
    run         = paragraph.add_run()
    run.add_picture(file_name, height=Inches(2.75))

    # Pareto
    paragraph = document.add_paragraph('Evolución del Riesgo')
    paragraph.style = document.styles['Intense Quote']    
    f_Pareto(df,file_name)
    paragraph   = document.add_paragraph() 
    run         = paragraph.add_run()
    run.add_picture(file_name, height=Inches(2.75))

    #from ST_AM_Contingencies_Ploty import Plot_Histogram_ENS
    #Plot_Histogram_ENS(df_mont,1,Years,Type=None,Factor='GWh')



    for year in  Years:
        # Add title 
        #asset     = assets.Asset_Portfolio[asset_id]
        paragraph = document.add_paragraph('Riesgo estimado para el año '+str(year))
        paragraph.style = document.styles['Title']

        # Histogram
        paragraph = document.add_paragraph('Histograma de Criticidad')
        paragraph.style = document.styles['Intense Quote']    

        f_Cr_Hist(df_cr_time[year],file_name,Factor='GWh')
        paragraph     = document.add_paragraph()  
        run = paragraph.add_run()
        run.add_picture(file_name, height=Inches(2.4))

        # Risk matrix ENS
        paragraph = document.add_paragraph('Matriz de riesgos')
        paragraph.style = document.styles['Intense Quote']            
        query         = 'Year=='+str(year)
        l_df          = df.query(query)
        f_Risk_Matrix(l_df,file_name,Factor='GWh') 
        paragraph     = document.add_paragraph()  
        run = paragraph.add_run()
        run.add_picture(file_name, height=Inches(2.4))

        # Risk matrix ENS
        #f_Risk_Matrix(l_df,file_name,Factor='h',y='SAIDI') 
        paragraph = document.add_paragraph('Priorización')
        paragraph.style = document.styles['Intense Quote']    
        f_Pareto_Year(l_df,file_name)
        paragraph     = document.add_paragraph()  
        run = paragraph.add_run()
        run.add_picture(file_name, height=Inches(2.4))





        # Table with indexes 
        '''template = document.tables[0]
        tbl      = template._tbl
        new_tbl = deepcopy(tbl)
        paragraph._p.addnext(new_tbl)

        document.tables[n_table]._cells[0].paragraphs[0].text = asset.type
        document.tables[n_table]._cells[0].paragraphs[1].text = 'Tipo de activo'

        document.tables[n_table]._cells[1].paragraphs[0].text = asset.data['Location']
        document.tables[n_table]._cells[1].paragraphs[1].text = 'Subestación'

        document.tables[n_table]._cells[2].paragraphs[0].text = str(asset.data['Serial'])
        document.tables[n_table]._cells[2].paragraphs[1].text = 'Serial'

        document.tables[n_table]._cells[3].paragraphs[0].text = asset.data['UC']
        document.tables[n_table]._cells[3].paragraphs[1].text = 'UC'

        n_table +=1 
# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #  
        paragraph = document.add_paragraph()  

        file_name = 'RESULTS/X1_REPORTS/X1_Images/fig.png'
        plot_historic_condition_by_asset(asset,file_name)
        run = paragraph.add_run()
        run.add_picture(file_name, height=Inches(3))

# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #
        new_tbl = deepcopy(tbl)
        paragraph._p.addnext(new_tbl)
        document.tables[n_table]._cells[3].paragraphs[0].text = str(round(asset.mttr,2))
        document.tables[n_table]._cells[3].paragraphs[1].text = 'MTTR - [h]'

        document.tables[n_table]._cells[1].paragraphs[0].text = str(round(asset.elap_life,2))
        document.tables[n_table]._cells[1].paragraphs[1].text = 'EL'

        if asset.hi == None:
            document.tables[n_table]._cells[0].paragraphs[0].text = 'None'
        else:
            document.tables[n_table]._cells[0].paragraphs[0].text = str(round(asset.hi,2))
        document.tables[n_table]._cells[0].paragraphs[1].text = 'HI '

        if asset.re==None:
            document.tables[n_table]._cells[2].paragraphs[0].text = 'None'
        else:
            document.tables[n_table]._cells[2].paragraphs[0].text = str(round(asset.re,2))
        document.tables[n_table]._cells[2].paragraphs[1].text = 'RE'
        n_table +=1 
# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #  
        paragraph = document.add_paragraph()  
        #file_name = 'RESULTS/X1_REPORTS/X1_Images/fig.png'
        Radar_Plot_by_Asset(asset,[report_date],file_name)
        run = paragraph.add_run()
        run.add_picture(file_name, height=Inches(3.5))

# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #
        new_tbl = deepcopy(tbl)
        paragraph._p.addnext(new_tbl)

        document.tables[n_table]._cells[0].paragraphs[0].text = asset.data['ID']
        document.tables[n_table]._cells[0].paragraphs[1].text = 'ID'

        document.tables[n_table]._cells[1].paragraphs[0].text = str(asset.data['Opt_Year'].year)
        document.tables[n_table]._cells[1].paragraphs[1].text = 'Año de ent. operación'

        document.tables[n_table]._cells[2].paragraphs[0].text = asset.data['Brand']
        document.tables[n_table]._cells[2].paragraphs[1].text = 'Marca'

        document.tables[n_table]._cells[3].paragraphs[0].text = asset.data['Model']
        document.tables[n_table]._cells[3].paragraphs[1].text = 'Modelo'
        n_table +=1''' 

        #document.add_page_break()
        
    #Save document 
    document.save('RESULTS/X1_REPORTS/Criticality.docx')